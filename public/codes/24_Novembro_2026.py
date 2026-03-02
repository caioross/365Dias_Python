"""
Script para gerar um contrato de aluguel preenchido com dados de locador e locatário.
Utiliza a biblioteca python-docx para manipulação de arquivos .docx.
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def preencher_contrato(modelo_path, dados_locador, dados_locatario, output_path):
    """
    Preenche um modelo de contrato de aluguel com os dados fornecidos.

    :param modelo_path: Caminho para o arquivo .docx do modelo do contrato.
    :param dados_locador: Dicionário contendo os dados do locador.
    :param dados_locatario: Dicionário contendo os dados do locatário.
    :param output_path: Caminho onde o contrato preenchido será salvo.
    """
    # Carrega o modelo de contrato
    doc = Document(modelo_path)

    # Substitui os placeholders pelos dados do locador
    for paragraph in doc.paragraphs:
        for key, value in dados_locador.items():
            if f'{{locador_{key}}}' in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{locador_{key}}}', value)

    # Substitui os placeholders pelos dados do locatário
    for paragraph in doc.paragraphs:
        for key, value in dados_locatario.items():
            if f'{{locatario_{key}}}' in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{locatario_{key}}}', value)

    # Salva o contrato preenchido
    doc.save(output_path)

def main():
    """
    Função principal que define os dados e chama a função para preencher o contrato.
    """
    # Dados do locador
    dados_locador = {
        'nome': 'João Silva',
        'endereco': 'Rua A, 123, Bairro X',
        'cpf': '123.456.789-00',
        'telefone': '(11) 98765-4321'
    }

    # Dados do locatário
    dados_locatario = {
        'nome': 'Maria Oliveira',
        'endereco': 'Rua B, 456, Bairro Y',
        'cpf': '987.654.321-00',
        'telefone': '(11) 12345-6789'
    }

    # Caminho para o modelo de contrato
    modelo_path = 'modelo_contrato.docx'

    # Caminho para salvar o contrato preenchido
    output_path = 'contrato_preenchido.docx'

    # Preenche o contrato
    preencher_contrato(modelo_path, dados_locador, dados_locatario, output_path)

if __name__ == '__main__':
    main()